"""Text extraction from uploaded documents.

Each supported type has a dedicated extractor. Images do not contain extractable
text without OCR (out of scope for this sprint), so they return an empty string
and rely on filename/context for classification.
"""
from __future__ import annotations

import io

from app.core.logging import get_logger

logger = get_logger(__name__)


def extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - a bad page shouldn't fail the whole doc
            logger.warning("Failed to extract a PDF page; skipping")
    return "\n".join(parts).strip()


def extract_docx(data: bytes) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    # Include table cell text as well.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs).strip()


def extract_image(data: bytes) -> str:
    # OCR is out of scope; validate the image is readable and return no text.
    from PIL import Image

    try:
        Image.open(io.BytesIO(data)).verify()
    except Exception:  # noqa: BLE001
        logger.warning("Uploaded image could not be verified")
    return ""


def extract_text(data: bytes, content_type: str) -> str:
    """Dispatch to the correct extractor based on content type."""
    if content_type == "application/pdf":
        return extract_pdf(data)
    if content_type == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return extract_docx(data)
    if content_type in ("image/png", "image/jpeg"):
        return extract_image(data)
    return ""
